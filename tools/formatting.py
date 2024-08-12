from typing import Any, Dict, List
from docx import Document


class JSONFormatting:
    @staticmethod
    def stringify(
            obj: dict or list,
            indent: int = 4,
            kv_connector: str = ":",
            ln_connector: str = "\n"
    ) -> str:
        """
        Convert a JSON object into a human-readable text format.

        Args:
            obj (dict or list): The JSON object to convert.
            indent (int): The current indentation level (used for recursive calls).
            kv_connector (str): The connector to use to connect a key to a value.
            ln_connector (str): The connector to use to connect lines.

        Returns:
            str: The human-readable text representation of the JSON object.
        """
        if not isinstance(obj, (dict, list)):
            return str(obj)

        text = ""
        indent_str = " " * (indent * 2)
        if isinstance(obj, dict):
            for key, value in obj.items():
                if isinstance(value, (dict, list)):
                    text += f"{indent_str}{key}{kv_connector} {ln_connector}"
                    text += JSONFormatting.stringify(value, indent + 1, kv_connector, ln_connector)
                elif value not in [None, "N/A"]:
                    text += f"{indent_str}{key}{kv_connector} {str(value)}{ln_connector}"
        elif isinstance(obj, list):
            for item in obj:
                text += indent_str + "- " + JSONFormatting.stringify(item, indent + 1, kv_connector, ln_connector) + ln_connector

        return text

    @staticmethod
    def clean(obj: dict or list):
        """
        Remove keys with null values (None or "N/A") from a JSON object.

        Args:
            obj (dict or list): The JSON object to clean.

        Returns:
            dict or list: The cleaned JSON object.
        """
        if isinstance(obj, dict):
            cleaned_dict = {}
            for key, value in obj.items():
                if value not in [None, "N/A", ""]:
                    if isinstance(value, (dict, list)):
                        cleaned_value = JSONFormatting.clean(value)
                        if cleaned_value:  # Only add to cleaned_dict if cleaned_value is not empty
                            cleaned_dict[key] = cleaned_value
                    else:
                        cleaned_dict[key] = value
            return cleaned_dict
        elif isinstance(obj, list):
            cleaned_list = []
            for item in obj:
                cleaned_item = JSONFormatting.clean(item)
                if cleaned_item:  # Only add to cleaned_list if cleaned_item is not empty
                    cleaned_list.append(cleaned_item)
            return cleaned_list
        else:
            return obj


class TextFormatting:
    @staticmethod
    def justify(text: str, length_thr: int = 120) -> str:
        """
        Print text with justified line length

        Args:
            text: text to print
            length_thr: the threshold length for each line

        Returns:
            str: justified text
        """
        justified = ""
        lines = text.split('\n')

        for line in lines:
            words = line.split()
            current_line = ""
            for word in words:
                if len(current_line) + len(word) + 1 > length_thr:
                    justified += current_line + "\n"
                    current_line = word + " "
                else:
                    current_line += word + " "

            if current_line:
                justified += current_line.rstrip() + "\n"
            else:
                justified += "\n"

        return justified


class DocxFormatting:
    def __init__(self, doc: Document) -> None:
        self.doc = doc

    class KeyChanger:
        def __init__(self, p, key, value) -> None:
            self.p = p
            self.key = key
            self.value = value
            self.run_text = ""
            self.runs_indexes: List = []
            self.run_char_indexes: List = []
            self.runs_to_change: Dict = {}

        def replace_key(self) -> None:
            run_index = 0
            for run in self.p.runs:
                self.run_text += run.text
                self.runs_indexes += [run_index for _ in run.text]
                self.run_char_indexes += [char_index for char_index, char in enumerate(run.text)]
                run_index += 1

            parsed_key_length = len(self.key)
            index_to_replace = self.run_text.find(self.key)

            for i in range(parsed_key_length):
                index = index_to_replace + i
                run_index = self.runs_indexes[index]
                run = self.p.runs[run_index]
                run_char_index = self.run_char_indexes[index]

                if not self.runs_to_change.get(run_index):
                    self.runs_to_change[run_index] = [char for char_index, char in enumerate(run.text)]

                run_to_change: Dict = self.runs_to_change.get(run_index)  # type: ignore[assignment]
                if index == index_to_replace:
                    run_to_change[run_char_index] = self.value
                else:
                    run_to_change[run_char_index] = ""

            # make the real replace
            for index, text in self.runs_to_change.items():
                run = self.p.runs[index]
                run.text = "".join(text)

    class Paragraph:
        def __init__(self, p) -> None:
            self.p = p

        @staticmethod
        def get_all(doc: Document) -> List[Any]:
            paragraphs = list()
            paragraphs.extend(DocxFormatting.Paragraph._get_paragraphs(doc))

            for section in doc.sections:
                paragraphs.extend(DocxFormatting.Paragraph._get_paragraphs(section.header))
                paragraphs.extend(DocxFormatting.Paragraph._get_paragraphs(section.footer))

            return paragraphs

        @staticmethod
        def _get_paragraphs(item: Any) -> Any:
            yield from item.paragraphs

            # get paragraphs from tables
            for table in item.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield paragraph

        def replace_key(self, key, value) -> None:
            if key in self.p.text:
                self._simple_replace_key(key, value)
                if key in self.p.text:
                    self._complex_replace_key(key, value)

        def _simple_replace_key(self, key, value) -> None:
            # try to replace a key in the paragraph runs, simpler alternative
            for run in self.p.runs:
                if key in run.text:
                    run.text = run.text.replace(key, value)

        def _complex_replace_key(self, key, value) -> None:
            # complex alternative, which check all broken items inside the runs
            while key in self.p.text:
                # if the key appears more than once in the paragraph, it will be replaced all
                key_changer = DocxFormatting.KeyChanger(self.p, key, value)
                key_changer.replace_key()

    def replace_keys(self, **kwargs: str) -> None:
        """
        Replace all the keys in the MS Word document with the values in the kwargs
        The required format for the keys inside the Word document is: ${key}

        Example usage:
            Word content = "Hello ${name}, your phone is ${phone}, is that okay?"

            doc = Document("document.docx")  # python-docx dependency

            docx_replacer = DocxKeyReplacer(doc)
            docx_replacer.replace_keys(name="Ivan", phone="+55123456789")
        """
        for key, value in kwargs.items():
            key = f"${{{key}}}"
            for p in self.Paragraph.get_all(self.doc):
                paragraph = self.Paragraph(p)
                paragraph.replace_key(key, str(value))
